"""Programmatic test fixture generation — no binary files committed."""

from __future__ import annotations

from pathlib import Path

import pymupdf
import pytest
from docx import Document as DocxDocument


@pytest.fixture
def tmp_fixtures(tmp_path: Path) -> Path:
    """Create a temporary directory with sample files for testing."""
    return tmp_path


# ── Text Fixtures ─────────────────────────────────────────────────────────────


@pytest.fixture
def txt_simple(tmp_path: Path) -> Path:
    p = tmp_path / "simple.txt"
    p.write_text(
        "This is paragraph one. It has two sentences.\n\n"
        "This is paragraph two.\n\n"
        "This is paragraph three with some longer content that spans more words.",
        encoding="utf-8",
    )
    return p


@pytest.fixture
def txt_unicode(tmp_path: Path) -> Path:
    p = tmp_path / "unicode.txt"
    p.write_text("Caf\u00e9 na\u00efve r\u00e9sum\u00e9.\n\n\u3053\u3093\u306b\u3061\u306f\u4e16\u754c\u3002", encoding="utf-8")
    return p


@pytest.fixture
def txt_latin1(tmp_path: Path) -> Path:
    p = tmp_path / "latin1.txt"
    p.write_bytes("Caf\xe9 na\xefve r\xe9sum\xe9.".encode("latin-1"))
    return p


@pytest.fixture
def txt_empty(tmp_path: Path) -> Path:
    p = tmp_path / "empty.txt"
    p.write_text("", encoding="utf-8")
    return p


@pytest.fixture
def txt_crlf(tmp_path: Path) -> Path:
    p = tmp_path / "crlf.txt"
    p.write_bytes(b"Line one.\r\n\r\nLine two.\r\n")
    return p


# ── Markdown Fixtures ─────────────────────────────────────────────────────────


@pytest.fixture
def md_full(tmp_path: Path) -> Path:
    p = tmp_path / "full.md"
    p.write_text(
        "# Main Title\n\n"
        "## Section One\n\n"
        "This is a paragraph with **bold** and *italic* text.\n\n"
        "- Item one\n"
        "- Item two\n"
        "- Item three\n\n"
        "1. First\n"
        "2. Second\n\n"
        "> This is a blockquote.\n\n"
        "---\n\n"
        "```python\nprint('hello')\n```\n\n"
        "A [link](https://example.com) and an ![image](pic.png).\n\n"
        "Some <b>HTML</b> and ~~strikethrough~~ text.\n\n"
        "Inline `code` here.\n",
        encoding="utf-8",
    )
    return p


@pytest.fixture
def md_headings_only(tmp_path: Path) -> Path:
    p = tmp_path / "headings.md"
    p.write_text("# H1\n\n## H2\n\n### H3\n\n#### H4\n", encoding="utf-8")
    return p


# ── DOCX Fixtures ─────────────────────────────────────────────────────────────


@pytest.fixture
def docx_simple(tmp_path: Path) -> Path:
    p = tmp_path / "simple.docx"
    doc = DocxDocument()
    doc.core_properties.title = "Test Document"
    doc.core_properties.author = "Test Author"
    doc.add_heading("Document Title", level=1)
    doc.add_paragraph("First paragraph of the document.")
    doc.add_heading("Section Two", level=2)
    doc.add_paragraph("Second paragraph with more content.")
    doc.save(str(p))
    return p


@pytest.fixture
def docx_with_table(tmp_path: Path) -> Path:
    p = tmp_path / "table.docx"
    doc = DocxDocument()
    doc.add_paragraph("Before the table.")
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(0, 2).text = "C1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"
    table.cell(1, 2).text = "C2"
    doc.add_paragraph("After the table.")
    doc.save(str(p))
    return p


@pytest.fixture
def docx_empty(tmp_path: Path) -> Path:
    p = tmp_path / "empty.docx"
    doc = DocxDocument()
    doc.save(str(p))
    return p


# ── PDF Fixtures ──────────────────────────────────────────────────────────────


@pytest.fixture
def pdf_simple(tmp_path: Path) -> Path:
    p = tmp_path / "simple.pdf"
    doc = pymupdf.open()
    page = doc.new_page()
    assert page is not None
    page.insert_text((72, 72), "Simple PDF Document", fontsize=20)
    page.insert_text((72, 120), "This is the first paragraph of text in the PDF.")
    page.insert_text((72, 145), "This is the second paragraph.")
    doc.save(str(p))
    doc.close()
    return p


@pytest.fixture
def pdf_multipage(tmp_path: Path) -> Path:
    p = tmp_path / "multipage.pdf"
    doc = pymupdf.open()
    for i in range(3):
        page = doc.new_page()
        assert page is not None
        page.insert_text((72, 72), f"Page {i + 1} content.")
    doc.save(str(p))
    doc.close()
    return p


@pytest.fixture
def pdf_empty_page(tmp_path: Path) -> Path:
    p = tmp_path / "empty_page.pdf"
    doc = pymupdf.open()
    doc.new_page()  # empty page — no text
    page2 = doc.new_page()
    assert page2 is not None
    page2.insert_text((72, 72), "Page two has text.")
    doc.save(str(p))
    doc.close()
    return p


@pytest.fixture
def pdf_with_metadata(tmp_path: Path) -> Path:
    p = tmp_path / "metadata.pdf"
    doc = pymupdf.open()
    page = doc.new_page()
    assert page is not None
    page.insert_text((72, 72), "Content here.")
    doc.set_metadata({"title": "My PDF Title", "author": "Jane Doe"})
    doc.save(str(p))
    doc.close()
    return p
